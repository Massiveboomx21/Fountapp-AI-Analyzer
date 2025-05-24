"""
Document Processor για AI Document Analyzer
"""
import PyPDF2
from docx import Document
import pytesseract
from PIL import Image
from pathlib import Path
from typing import List, Dict, Optional
import re
from config import config
from utils.logger import setup_logger

logger = setup_logger()

class DocumentProcessor:
    """Επεξεργασία και εξαγωγή κειμένου από έγγραφα"""
    
    def __init__(self):
        self.max_chunk_size = config.MAX_CHUNK_SIZE
        
    def process_document(self, file_info: Dict) -> Dict:
        """
        Επεξεργασία εγγράφου και εξαγωγή κειμένου
        
        Args:
            file_info: Πληροφορίες αρχείου από FileScanner
            
        Returns:
            Dict με extracted text και metadata
        """
        filepath = file_info['filepath']
        file_extension = file_info['file_extension']
        
        logger.info(f"Επεξεργασία αρχείου: {file_info['filename']}")
        
        try:
            # Εξαγωγή κειμένου ανάλογα με τον τύπο αρχείου
            if file_extension == '.pdf':
                text = self._extract_from_pdf(filepath)
            elif file_extension == '.docx':
                text = self._extract_from_docx(filepath)
            elif file_extension == '.txt':
                text = self._extract_from_txt(filepath)
            elif file_extension in ['.png', '.jpg', '.jpeg']:
                text = self._extract_from_image(filepath)
            else:
                raise ValueError(f"Μη υποστηριζόμενος τύπος αρχείου: {file_extension}")
            
            # Καθαρισμός κειμένου
            cleaned_text = self._clean_text(text)
            
            # Διαίρεση σε chunks
            chunks = self._split_into_chunks(cleaned_text)
            
            # Metadata εξαγωγής
            extraction_meta = {
                'original_length': len(text),
                'cleaned_length': len(cleaned_text),
                'chunk_count': len(chunks),
                'word_count': len(cleaned_text.split()),
                'extraction_method': file_extension[1:]  # Αφαίρεση της τελείας
            }
            
            result = {
                'success': True,
                'original_text': text,
                'cleaned_text': cleaned_text,
                'chunks': chunks,
                'metadata': extraction_meta,
                'error': None
            }
            
            logger.info(f"Επιτυχής επεξεργασία: {file_info['filename']} "
                       f"({extraction_meta['word_count']} λέξεις, {len(chunks)} chunks)")
            
            return result
            
        except Exception as e:
            logger.error(f"Σφάλμα επεξεργασίας {file_info['filename']}: {str(e)}")
            return {
                'success': False,
                'original_text': '',
                'cleaned_text': '',
                'chunks': [],
                'metadata': {},
                'error': str(e)
            }
    
    def _extract_from_pdf(self, filepath: str) -> str:
        """Εξαγωγή κειμένου από PDF"""
        text = ""
        try:
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Σελίδα {page_num + 1} ---\n"
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Σφάλμα ανάγνωσης σελίδας {page_num + 1}: {e}")
                        continue
                
                if not text.strip():
                    logger.warning(f"Δεν βρέθηκε κείμενο στο PDF: {filepath}")
                    
        except Exception as e:
            raise Exception(f"Σφάλμα ανάγνωσης PDF: {e}")
        
        return text
    
    def _extract_from_docx(self, filepath: str) -> str:
        """Εξαγωγή κειμένου από DOCX"""
        try:
            doc = Document(filepath)
            
            # Εξαγωγή κειμένου από παραγράφους
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Εξαγωγή κειμένου από πίνακες
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Συνδυασμός όλου του κειμένου
            all_text = []
            if paragraphs:
                all_text.append("\n".join(paragraphs))
            if tables_text:
                all_text.append("\n--- Πίνακες ---\n")
                all_text.append("\n".join(tables_text))
            
            return "\n\n".join(all_text)
            
        except Exception as e:
            raise Exception(f"Σφάλμα ανάγνωσης DOCX: {e}")
    
    def _extract_from_txt(self, filepath: str) -> str:
        """Εξαγωγή κειμένου από TXT"""
        try:
            # Δοκιμή διαφορετικών encodings
            encodings = ['utf-8', 'utf-8-sig', 'cp1253', 'iso-8859-7', 'latin1']
            
            for encoding in encodings:
                try:
                    with open(filepath, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content:  # Αν διαβάστηκε επιτυχώς
                            logger.debug(f"Χρήση encoding {encoding} για {filepath}")
                            return content
                except UnicodeDecodeError:
                    continue
            
            # Αν δεν δούλεψε κανένα encoding
            raise Exception("Δεν μπόρεσα να ανοίξω το αρχείο με κανένα γνωστό encoding")
            
        except Exception as e:
            raise Exception(f"Σφάλμα ανάγνωσης TXT: {e}")
    
    def _extract_from_image(self, filepath: str) -> str:
        """Εξαγωγή κειμένου από εικόνα με OCR"""
        try:
            # Άνοιγμα εικόνας
            image = Image.open(filepath)
            
            # OCR με υποστήριξη ελληνικών και αγγλικών
            text = pytesseract.image_to_string(
                image, 
                lang='ell+eng',
                config='--psm 6'  # Uniform text block
            )
            
            if not text.strip():
                logger.warning(f"Δεν βρέθηκε κείμενο στην εικόνα: {filepath}")
            
            return text
            
        except Exception as e:
            raise Exception(f"Σφάλμα OCR: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Καθαρισμός κειμένου"""
        if not text:
            return ""
        
        # Αντικατάσταση πολλαπλών κενών διαστημάτων
        text = re.sub(r'\s+', ' ', text)
        
        # Αφαίρεση πολλαπλών line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Καθαρισμός ειδικών χαρακτήρων που μπορεί να προκαλούν προβλήματα
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Trim και αφαίρεση κενών γραμμών στην αρχή/τέλος
        text = text.strip()
        
        return text
    
    def _split_into_chunks(self, text: str) -> List[str]:
        """Διαίρεση κειμένου σε chunks για AI processing"""
        if not text:
            return []
        
        chunks = []
        current_chunk = ""
        
        # Διαίρεση σε παραγράφους πρώτα
        paragraphs = text.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # Αν η παράγραφος χωράει στο τρέχον chunk
            if len(current_chunk + paragraph) <= self.max_chunk_size:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                # Αν υπάρχει τρέχον chunk, προσθήκη στη λίστα
                if current_chunk:
                    chunks.append(current_chunk)
                
                # Αν η παράγραφος είναι πολύ μεγάλη, διαίρεση σε προτάσεις
                if len(paragraph) > self.max_chunk_size:
                    sentence_chunks = self._split_long_paragraph(paragraph)
                    chunks.extend(sentence_chunks)
                    current_chunk = ""
                else:
                    current_chunk = paragraph
        
        # Προσθήκη του τελευταίου chunk
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _split_long_paragraph(self, paragraph: str) -> List[str]:
        """Διαίρεση μεγάλης παραγράφου σε μικρότερα chunks"""
        chunks = []
        
        # Διαίρεση σε προτάσεις
        sentences = re.split(r'[.!?]+', paragraph)
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Αν η πρόταση χωράει στο τρέχον chunk
            if len(current_chunk + sentence) <= self.max_chunk_size:
                if current_chunk:
                    current_chunk += ". " + sentence
                else:
                    current_chunk = sentence
            else:
                # Αποθήκευση τρέχοντος chunk και έναρξη νέου
                if current_chunk:
                    chunks.append(current_chunk + ".")
                
                # Αν η πρόταση είναι πολύ μεγάλη, διαίρεση σε λέξεις
                if len(sentence) > self.max_chunk_size:
                    word_chunks = self._split_by_words(sentence)
                    chunks.extend(word_chunks)
                    current_chunk = ""
                else:
                    current_chunk = sentence
        
        # Προσθήκη τελευταίου chunk
        if current_chunk:
            chunks.append(current_chunk + ".")
        
        return chunks
    
    def _split_by_words(self, text: str) -> List[str]:
        """Διαίρεση κειμένου σε chunks βάσει λέξεων (τελευταία επιλογή)"""
        words = text.split()
        chunks = []
        current_chunk = ""
        
        for word in words:
            if len(current_chunk + " " + word) <= self.max_chunk_size:
                if current_chunk:
                    current_chunk += " " + word
                else:
                    current_chunk = word
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = word
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def get_document_preview(self, text: str, max_length: int = 500) -> str:
        """Δημιουργία preview εγγράφου"""
        if len(text) <= max_length:
            return text
        
        # Κοπή στο πλησιέστερο space για να μην κόψουμε λέξη
        preview = text[:max_length]
        last_space = preview.rfind(' ')
        
        if last_space > max_length * 0.8:  # Αν το space είναι αρκετά κοντά
            preview = preview[:last_space]
        
        return preview + "..."
    
    def validate_extracted_text(self, text: str) -> Dict:
        """Επικύρωση εξαγόμενου κειμένου"""
        validation = {
            'is_valid': False,
            'has_content': False,
            'estimated_language': 'unknown',
            'quality_score': 0.0,
            'warnings': []
        }
        
        if not text or not text.strip():
            validation['warnings'].append("Δεν βρέθηκε κείμενο")
            return validation
        
        validation['has_content'] = True
        
        # Εκτίμηση ποιότητας κειμένου
        word_count = len(text.split())
        char_count = len(text)
        
        if word_count < 5:
            validation['warnings'].append("Πολύ λίγες λέξεις")
            validation['quality_score'] = 0.3
        elif char_count / word_count > 20:  # Πολύ μεγάλες "λέξεις" - πιθανώς σφάλμα OCR
            validation['warnings'].append("Πιθανό σφάλμα εξαγωγής κειμένου")
            validation['quality_score'] = 0.5
        else:
            validation['quality_score'] = min(1.0, word_count / 100)
        
        # Εκτίμηση γλώσσας (απλοποιημένη)
        greek_chars = sum(1 for c in text if '\u0370' <= c <= '\u03FF' or '\u1F00' <= c <= '\u1FFF')
        latin_chars = sum(1 for c in text if c.isalpha() and c.isascii())
        
        if greek_chars > latin_chars:
            validation['estimated_language'] = 'greek'
        elif latin_chars > 0:
            validation['estimated_language'] = 'english'
        
        validation['is_valid'] = validation['has_content'] and validation['quality_score'] > 0.2
        
        return validation
